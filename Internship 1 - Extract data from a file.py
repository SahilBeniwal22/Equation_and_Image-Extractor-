"""This program extracts all mathematical equation and diagrams or graphs that are in image format"""
"""In this program os module is used to open the docx file
re module is used to find the equations in the data 
rest modules are used to perform operations to the docx files"""

from docx import Document
import re
import docx2txt as d2t
import os

file = Document()
file.add_paragraph("Welcome to Internship code made by :- SAHIL BENIWAL \n"
                   "Here are the Equations Extracted from your data.")

# This functions extracts the images from the docx file
def image_extractor(file_path,image_path,get_text=False):
    text=d2t.process(file_path,image_path)

# This function extracts the text and gives to equation extractor
def text_extractor(file_path,get_text=True):
    print("ALL MATHEMATICAL EQUATIONS IN YOUR GIVEN DATA ARE LISTED BELOW : -")
    text=d2t.process(file_path)
    if (get_text):
        return text

# This function extracts the equations and then add them to a separate file
def eqn_extractor(data):
    save_path="G:\Sahil\equation.docx"
    # print("THE EQUATION ARE :-")
    if(re.findall(r"[0-9]+",item)):
        newitem=item
        # print(item,end=""+",")
        if(re.findall(r"[+,-]+",newitem)):
            print(newitem,end=""+",")
            file.add_paragraph(newitem)
            file.save(save_path)
            # os.system("start G:\Sahil\equation.docx")
# You can uncomment above line if you want to open the saved docx file.

    else:
        print(end="")

if __name__ == '__main__':
    path="G:/Sahil/Python Codes/Internship/data.docx"
    image_path="G:\Sahil\Python Codes\Internship\Images and graphs"

    # You can change the path according to your own requirement

    image_extractor(path,image_path)
    data=text_extractor(path)
item = " "
for word in data:
    if (word != " "):
        item = item + word
    else:
        eqn_extractor(item)
        item = " "
